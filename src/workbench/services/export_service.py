"""Export pipeline (P6): canonical manuscript (DB) → Markdown, LaTeX, HTML, DOCX,
BibTeX, and a provenance manifest with checksums.

Rules:
- The DB is the canonical source; exports are derived artifacts, never edited in place.
- Every exported claim keeps its support state visible ([support: ...] annotations in
  draft formats); export never launders uncertainty.
- The manifest records source access levels, AI-provenance summary, audit findings at
  export time, and sha256 of every emitted file. Export implies nothing was submitted
  or published anywhere.
- PDF is intentionally deferred (WeasyPrint/GTK on Windows); LaTeX output compiles with
  any standard toolchain.
"""

import hashlib
import json
import re
from datetime import UTC, datetime
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.orm import Session

from ..config import get_settings
from ..models import Claim, ClaimEvidence, Excerpt, ResearchObject, Source
from ..vocab import ObjectKind
from . import audits, authoring, research


def _latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
        "_": r"\_", "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in text)


def _bib_key(source: Source) -> str:
    first_author = (source.authors.split(";")[0].split(",")[0].strip() or "anon").lower()
    first_author = re.sub(r"[^a-z0-9]", "", first_author) or "anon"
    return f"{first_author}{source.year or 'nd'}"


def _collect(session: Session, manuscript_id: str):
    manuscript = session.get(ResearchObject, manuscript_id)
    if manuscript is None or manuscript.kind != ObjectKind.MANUSCRIPT:
        raise research.IntegrityError("manuscript not found")
    sections = authoring.manuscript_sections(session, manuscript_id)
    claims: dict[str, Claim] = {}
    sources: dict[str, Source] = {}
    for section in sections:
        for cid in section.body.get("claim_ids", []):
            claim = session.get(Claim, cid)
            if claim is None:
                continue
            claims[cid] = claim
            for ev in session.scalars(
                select(ClaimEvidence).where(ClaimEvidence.claim_id == cid)
            ):
                if ev.excerpt_id:
                    excerpt = session.get(Excerpt, ev.excerpt_id)
                    if excerpt:
                        src = session.get(Source, excerpt.source_id)
                        if src:
                            sources[src.id] = src
    return manuscript, sections, claims, sources


def _claims_block(claim_ids: list[str], claims: dict[str, Claim], sources: dict[str, Source],
                  session: Session, fmt: str) -> list[str]:
    lines = []
    for cid in claim_ids:
        claim = claims.get(cid)
        if claim is None:
            continue
        cites = []
        for ev in session.scalars(select(ClaimEvidence).where(ClaimEvidence.claim_id == cid)):
            if ev.excerpt_id:
                excerpt = session.get(Excerpt, ev.excerpt_id)
                src = sources.get(excerpt.source_id) if excerpt else None
                if src:
                    cites.append((_bib_key(src), excerpt.locator))
        if fmt == "tex":
            cite = " ".join(f"\\cite{{{k}}}" for k, _loc in cites)
            lines.append(f"% [support: {claim.support}]\n{_latex_escape(claim.text)} {cite}".strip())
        else:
            cite = " ".join(f"[@{k}, {loc}]" for k, loc in cites)
            lines.append(f"{claim.text} {cite} *[support: {claim.support}]*".strip())
    return lines


_WEASYPRINT_CACHE: bool | None = None


def weasyprint_available() -> bool:
    """Probe whether WeasyPrint can actually be USED (import succeeds AND its native
    GTK/Pango libraries load). Catches OSError from missing DLLs, not just ImportError.
    Cached because the probe is not free."""
    global _WEASYPRINT_CACHE
    if _WEASYPRINT_CACHE is None:
        try:
            import weasyprint  # noqa: F401

            _WEASYPRINT_CACHE = True
        except Exception:  # ImportError, OSError (missing libgobject), etc.
            _WEASYPRINT_CACHE = False
    return _WEASYPRINT_CACHE


def _weasyprint_pdf(html: str) -> bytes:
    import weasyprint

    return weasyprint.HTML(string=html).write_pdf()


def _render_pdf(title: str, html: str, paragraphs: list[str], mode: str) -> tuple[bytes, str]:
    """Return (pdf_bytes, renderer_used). mode: auto|weasyprint|minimal.
    'weasyprint' requires it; 'auto' uses it when available and falls back; 'minimal'
    always uses the built-in deterministic renderer."""
    if mode == "minimal":
        return _minimal_pdf(title, paragraphs), "minimal"
    if mode == "weasyprint":
        if not weasyprint_available():
            raise research.IntegrityError(
                "pdf_renderer=weasyprint but WeasyPrint/GTK is unavailable on this system"
            )
        return _weasyprint_pdf(html), "weasyprint"
    # auto
    if weasyprint_available():
        return _weasyprint_pdf(html), "weasyprint"
    return _minimal_pdf(title, paragraphs), "minimal"


def _minimal_pdf(title: str, paragraphs: list[str]) -> bytes:
    """Dependency-free deterministic text PDF (concept from Nexus MinimalPdfRenderer):
    Helvetica, naive wrap, A4-ish pages. A faithful fallback — not typeset output; the
    LaTeX export is the publication-quality path."""

    def esc(text: str) -> str:
        return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")

    lines: list[str] = []
    for para in [title, ""] + paragraphs:
        for raw_line in para.splitlines() or [""]:
            words, current = raw_line.split(), ""
            if not words:
                lines.append("")
            while words:
                word = words.pop(0)
                if len(current) + len(word) + 1 > 90:
                    lines.append(current)
                    current = word
                else:
                    current = f"{current} {word}".strip()
            if current:
                lines.append(current)
        lines.append("")
    pages = [lines[i : i + 48] for i in range(0, len(lines), 48)] or [[]]

    objects: list[bytes] = []
    page_ids = [4 + i * 2 for i in range(len(pages))]
    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for pid, page in zip(page_ids, pages, strict=True):
        content = ["BT /F1 10 Tf 50 780 Td 14 TL"]
        for line in page:
            content.append(f"({esc(line)}) Tj T*")
        content.append("ET")
        stream = "\n".join(content).encode("latin-1", errors="replace")
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {pid + 1} 0 R >>".encode()
        )
        objects.append(
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        )
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for off in offsets[1:]:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()
    return bytes(out)


def _jats_xml(manuscript, sections, claims, sources, session: Session) -> str:
    """Minimal JATS 1.3-shaped article XML (front/body/back). Structural export for
    interchange; not validated against the full JATS DTD."""
    from xml.sax.saxutils import escape

    parts = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<article xmlns:xlink="http://www.w3.org/1999/xlink" article-type="research-article">',
        "<front><article-meta>",
        f"<title-group><article-title>{escape(manuscript.title)}</article-title></title-group>",
        "</article-meta></front>",
        "<body>",
    ]
    for s in sections:
        parts.append(f'<sec id="{s.id}"><title>{escape(s.title)}</title>')
        if s.body.get("text"):
            parts.append(f"<p>{escape(s.body['text'])}</p>")
        for line in _claims_block(s.body.get("claim_ids", []), claims, sources, session, "md"):
            parts.append(f"<p>{escape(line)}</p>")
        parts.append("</sec>")
    parts.append("</body><back><ref-list>")
    for src in sources.values():
        parts.append(
            f'<ref id="{_bib_key(src)}"><mixed-citation>{escape(src.authors)} '
            f"({src.year or 'n.d.'}). {escape(src.title)}. {escape(src.venue)}."
            + (f" doi:{escape(src.doi)}" if src.doi else "")
            + "</mixed-citation></ref>"
        )
    parts.append("</ref-list></back></article>")
    return "\n".join(parts)


def _export_supplements(session: Session, project_id: str, out_dir: Path) -> list[dict]:
    """Copy the project's figures/tables into the bundle with their data-provenance and a
    live staleness check, so a reviewer can trace every artifact back to its data."""
    from sqlalchemy import select

    from ..models import ResearchObject
    from ..vocab import ObjectKind
    from .figures import Dataset

    supp_dir = out_dir / "supplements"
    out: list[dict] = []
    for kind in (ObjectKind.FIGURE, ObjectKind.TABLE):
        for art in session.scalars(
            select(ResearchObject).where(
                ResearchObject.project_id == project_id, ResearchObject.kind == kind,
                ResearchObject.deleted_at.is_(None),
            )
        ):
            supp_dir.mkdir(parents=True, exist_ok=True)
            ds = session.get(ResearchObject, art.body.get("dataset_id"))
            stale = None
            if ds is not None:
                current = Dataset(ds.body["columns"], ds.body["rows"]).hash()
                stale = current != art.body.get("data_hash")
            entry = {"id": art.id, "kind": str(kind), "number": art.body.get("number"),
                     "title": art.title, "data_hash": art.body.get("data_hash"),
                     "stale": stale, "caption": art.body.get("caption")}
            if kind == ObjectKind.FIGURE:
                png = art.body.get("png_path")
                if png and Path(png).is_file():
                    dest = supp_dir / f"figure_{art.body.get('number')}.png"
                    dest.write_bytes(Path(png).read_bytes())
                    svg = art.body.get("svg_path")
                    if svg and Path(svg).is_file():
                        (supp_dir / f"figure_{art.body.get('number')}.svg").write_text(
                            Path(svg).read_text(encoding="utf-8"), encoding="utf-8")
                    entry["file"] = dest.name
            else:
                dest = supp_dir / f"table_{art.body.get('number')}.md"
                dest.write_text(art.body.get("markdown", ""), encoding="utf-8")
                entry["file"] = dest.name
            out.append(entry)
    return out


def export_manuscript(
    session: Session, manuscript_id: str, *, formats: list[str] | None = None
) -> dict:
    formats = formats or ["md", "tex", "html", "docx", "bib", "pdf", "jats"]
    manuscript, sections, claims, sources = _collect(session, manuscript_id)
    out_dir = Path(get_settings().data_dir) / "exports" / manuscript_id
    out_dir.mkdir(parents=True, exist_ok=True)
    written: dict[str, Path] = {}

    # --- Markdown (canonical draft rendering) ---
    if "md" in formats:
        md = [f"# {manuscript.title}", ""]
        for s in sections:
            md += [f"## {s.title}", ""]
            if s.body.get("text"):
                md += [s.body["text"], ""]
            block = _claims_block(s.body.get("claim_ids", []), claims, sources, session, "md")
            if block:
                md += ["**Claims:**", ""] + [f"- {line}" for line in block] + [""]
        if sources:
            md += ["## References", ""]
            md += [
                f"- `{_bib_key(s)}`: {s.authors} ({s.year or 'n.d.'}). {s.title}. "
                f"{s.venue}. {'doi:' + s.doi if s.doi else ''} [access: {s.access}]"
                for s in sources.values()
            ]
        written["md"] = out_dir / "manuscript.md"
        written["md"].write_text("\n".join(md), encoding="utf-8")

    # --- LaTeX ---
    if "tex" in formats:
        tex = [
            r"\documentclass{article}", r"\usepackage[utf8]{inputenc}", "",
            f"\\title{{{_latex_escape(manuscript.title)}}}",
            r"\begin{document}", r"\maketitle", "",
        ]
        for s in sections:
            tex.append(f"\\section{{{_latex_escape(s.title)}}}")
            if s.body.get("text"):
                tex.append(_latex_escape(s.body["text"]))
            tex += _claims_block(s.body.get("claim_ids", []), claims, sources, session, "tex")
            tex.append("")
        if sources:
            tex += [r"\bibliographystyle{plain}", r"\bibliography{references}"]
        tex.append(r"\end{document}")
        written["tex"] = out_dir / "manuscript.tex"
        written["tex"].write_text("\n".join(tex), encoding="utf-8")

    # --- HTML (deterministic, no external assets) ---
    if "html" in formats:
        import html as _html

        parts = [f"<h1>{_html.escape(manuscript.title)}</h1>"]
        for s in sections:
            parts.append(f"<h2>{_html.escape(s.title)}</h2>")
            if s.body.get("text"):
                parts.append(f"<p>{_html.escape(s.body['text'])}</p>")
            block = _claims_block(s.body.get("claim_ids", []), claims, sources, session, "md")
            if block:
                parts.append("<ul>" + "".join(f"<li>{_html.escape(b)}</li>" for b in block) + "</ul>")
        written["html"] = out_dir / "manuscript.html"
        written["html"].write_text(
            "<!doctype html><meta charset='utf-8'><title>"
            + manuscript.title + "</title>" + "\n".join(parts),
            encoding="utf-8",
        )

    # --- DOCX (python-docx; lazy import) ---
    if "docx" in formats:
        try:
            import docx
        except ImportError:
            docx = None
        if docx is not None:
            document = docx.Document()
            document.add_heading(manuscript.title, level=0)
            for s in sections:
                document.add_heading(s.title, level=1)
                if s.body.get("text"):
                    document.add_paragraph(s.body["text"])
                for line in _claims_block(
                    s.body.get("claim_ids", []), claims, sources, session, "md"
                ):
                    document.add_paragraph(line, style="List Bullet")
            written["docx"] = out_dir / "manuscript.docx"
            document.save(str(written["docx"]))

    # --- BibTeX ---
    if "bib" in formats and sources:
        entries = []
        for s in sources.values():
            authors = " and ".join(a.strip() for a in s.authors.split(";") if a.strip())
            fields = [f"  title = {{{s.title}}}", f"  author = {{{authors}}}"]
            if s.year:
                fields.append(f"  year = {{{s.year}}}")
            if s.venue:
                fields.append(f"  journal = {{{s.venue}}}")
            if s.doi:
                fields.append(f"  doi = {{{s.doi}}}")
            entries.append(f"@article{{{_bib_key(s)},\n" + ",\n".join(fields) + "\n}")
        written["bib"] = out_dir / "references.bib"
        written["bib"].write_text("\n\n".join(entries), encoding="utf-8")

    # --- PDF (WeasyPrint typeset when available, else deterministic fallback) ---
    extra_manifest: dict = {}
    if "pdf" in formats:
        paragraphs = []
        html_parts = [f"<h1>{__import__('html').escape(manuscript.title)}</h1>"]
        for s in sections:
            paragraphs.append(s.title.upper())
            html_parts.append(f"<h2>{__import__('html').escape(s.title)}</h2>")
            if s.body.get("text"):
                paragraphs.append(s.body["text"])
                html_parts.append(f"<p>{__import__('html').escape(s.body['text'])}</p>")
            block = _claims_block(s.body.get("claim_ids", []), claims, sources, session, "md")
            paragraphs.extend(block)
            if block:
                html_parts.append(
                    "<ul>" + "".join(f"<li>{__import__('html').escape(b)}</li>" for b in block)
                    + "</ul>"
                )
        pdf_html = (
            "<!doctype html><meta charset='utf-8'><title>"
            + __import__("html").escape(manuscript.title) + "</title>" + "\n".join(html_parts)
        )
        pdf_bytes, renderer_used = _render_pdf(
            manuscript.title, pdf_html, paragraphs, get_settings().pdf_renderer
        )
        written["pdf"] = out_dir / "manuscript.pdf"
        written["pdf"].write_bytes(pdf_bytes)
        extra_manifest["pdf_renderer"] = renderer_used

    # --- JATS XML (validated against bundled JATS-subset DTD when lxml present) ---
    if "jats" in formats:
        from .jats import validate_jats

        jats_text = _jats_xml(manuscript, sections, claims, sources, session)
        written["jats"] = out_dir / "manuscript.jats.xml"
        written["jats"].write_text(jats_text, encoding="utf-8")
        extra_manifest["jats_validation"] = validate_jats(jats_text).as_dict()

    # --- Supplements: figures & tables with data provenance ---
    supplements = _export_supplements(session, manuscript.project_id, out_dir)
    if supplements:
        extra_manifest["supplements"] = supplements

    # --- Provenance manifest ---
    findings = audits.audit_manuscript(session, manuscript_id)
    manifest = {
        "manuscript_id": manuscript_id,
        "title": manuscript.title,
        "exported_at": datetime.now(UTC).isoformat(),
        "exported_by": "paper-workbench 0.1.0 (export != submission/publication)",
        "sections": [s.id for s in sections],
        "claims": {
            cid: {"support": str(c.support), "text": c.text} for cid, c in claims.items()
        },
        "sources": {
            sid: {
                "bib_key": _bib_key(s), "title": s.title, "doi": s.doi,
                "access": str(s.access), "license": s.license,
                "human_verified": s.human_verified,
            }
            for sid, s in sources.items()
        },
        "audit_findings_at_export": findings,
        "files": {},
        **extra_manifest,
    }
    for fmt, path in written.items():
        manifest["files"][fmt] = {
            "path": path.name,
            "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        }
    manifest_path = out_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    return {
        "out_dir": str(out_dir),
        "files": {fmt: str(p) for fmt, p in written.items()} | {"manifest": str(manifest_path)},
        "audit_findings": len(findings),
        **extra_manifest,
    }
